"""
v28 STAGE 22 — multi-format file ingestion → the S21 grounding pipeline.
=========================================================================
Detect a file's format and route it to the right extractor; the extracted text/structure feeds S21
(grounding). Stdlib formats (json / ipynb / csv / txt / code / zip / tar) ALWAYS work; office/PDF/image
formats use optional libraries and degrade HONESTLY when those are unavailable.

★ HONEST (§1.1, §1.9, §5.10) ★: (1) extraction confidence is LABELED — a scanned PDF with no text layer is
EXTRACTED_LOWCONF "(needs OCR)", a missing library is [BLOCKED: install …], a corrupt file is FAILED. We
never fabricate text. (2) Optional imports are guarded with `except BaseException` — in THIS sandbox `pypdf`
pulls a broken `cryptography` rust backend that raises a *PanicException* (a BaseException), so PDF is
[BLOCKED] here while docx/xlsx work; that is reported truthfully. (3) Extraction accuracy is corpus-specific
(vendor numbers like Docling 97.9% are not universal); low-confidence output is labeled, not trusted.
"""
from __future__ import annotations

import csv as _csv
import io
import json
import tarfile
import zipfile
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple


def _try(modname: str) -> bool:
    # Probe an optional import. pypdf's broken rust backend panics to fd 2 (stderr) AND raises a
    # BaseException — so we redirect fd 2 to /dev/null during the probe to keep output clean, and catch
    # BaseException (not just ImportError). (file_ingest may use os; the no-os rule is claude_agent-only.)
    import os
    saved = os.dup(2)
    devnull = os.open(os.devnull, os.O_WRONLY)
    try:
        os.dup2(devnull, 2)
        __import__(modname)
        return True
    except BaseException:  # noqa: BLE001
        return False
    finally:
        os.dup2(saved, 2)
        os.close(devnull)
        os.close(saved)


HAVE = {"pdf": _try("pypdf"), "docx": _try("docx"), "xlsx": _try("openpyxl"), "ocr": _try("pytesseract")}
import shutil as _shutil
HAVE["ocr"] = HAVE["ocr"] and (_shutil.which("tesseract") is not None)   # OCR also needs the system binary


@dataclass
class IngestResult:
    fmt: str
    status: str                 # EXTRACTED | EXTRACTED_LOWCONF | BLOCKED | FAILED
    text: str = ""
    structured: object = None
    confidence: str = "high"    # high | low | none
    detail: str = ""
    members: List[str] = field(default_factory=list)

    def __str__(self):
        return f"{self.status} [{self.fmt}] conf={self.confidence} ({len(self.text)} chars) {self.detail}"


_MAGIC = [(b"%PDF", "pdf"), (b"PK\x03\x04", "zip"), (b"\x89PNG", "png"),
          (b"\xff\xd8\xff", "jpeg"), (b"GIF8", "gif"), (b"\x1f\x8b", "gzip")]


def detect_format(filename: str, data: bytes) -> str:
    """Format by EXTENSION first (docx/xlsx/ipynb are zip containers — extension disambiguates), then magic."""
    name = (filename or "").lower()
    for ext, fmt in ((".pdf", "pdf"), (".docx", "docx"), (".xlsx", "xlsx"), (".csv", "csv"),
                     (".json", "json"), (".ipynb", "ipynb"), (".zip", "zip"), (".tar", "tar"),
                     (".tar.gz", "tar"), (".tgz", "tar"), (".png", "image"), (".jpg", "image"),
                     (".jpeg", "image"), (".gif", "image"), (".md", "text"), (".txt", "text"),
                     (".py", "code"), (".go", "code"), (".c", "code"), (".rs", "code"), (".js", "code")):
        if name.endswith(ext):
            return fmt
    for magic, fmt in _MAGIC:
        if data.startswith(magic):
            return {"png": "image", "jpeg": "image", "gif": "image", "gzip": "tar"}.get(fmt, fmt)
    try:
        data.decode("utf-8")
        return "text"
    except UnicodeDecodeError:
        return "unknown"


# ── per-format extractors ───────────────────────────────────────────────────────────────────────────
def _ingest_json(data: bytes) -> IngestResult:
    obj = json.loads(data.decode("utf-8"))
    return IngestResult("json", "EXTRACTED", json.dumps(obj, indent=2)[:100000], obj)


def _ingest_ipynb(data: bytes) -> IngestResult:
    nb = json.loads(data.decode("utf-8"))
    cells = []
    for c in nb.get("cells", []):
        src = "".join(c.get("source", []))
        cells.append(f"# [{c.get('cell_type')}]\n{src}")
    return IngestResult("ipynb", "EXTRACTED", "\n\n".join(cells), nb,
                        detail=f"{len(nb.get('cells', []))} cells")


def _ingest_csv(data: bytes) -> IngestResult:
    rows = list(_csv.reader(io.StringIO(data.decode("utf-8"))))
    text = "\n".join(", ".join(r) for r in rows)
    return IngestResult("csv", "EXTRACTED", text, rows, detail=f"{len(rows)} rows")


def _ingest_text(data: bytes, fmt: str = "text") -> IngestResult:
    return IngestResult(fmt, "EXTRACTED", data.decode("utf-8", "replace"))


def _ingest_archive(data: bytes, fmt: str) -> IngestResult:
    members: List[str] = []
    texts: List[str] = []
    try:
        if fmt == "zip":
            zf = zipfile.ZipFile(io.BytesIO(data))
            names = zf.namelist()
            for nm in names:
                members.append(nm)
                if nm.endswith((".txt", ".md", ".py", ".json", ".csv", ".go", ".c", ".rs", ".js")):
                    sub = ingest(zf.read(nm), nm)
                    if sub.status.startswith("EXTRACTED"):
                        texts.append(f"# {nm}\n{sub.text}")
        else:
            tf = tarfile.open(fileobj=io.BytesIO(data))
            for ti in tf.getmembers():
                members.append(ti.name)
                if ti.isfile() and ti.name.endswith((".txt", ".md", ".py", ".json", ".csv")):
                    f = tf.extractfile(ti)
                    if f:
                        sub = ingest(f.read(), ti.name)
                        if sub.status.startswith("EXTRACTED"):
                            texts.append(f"# {ti.name}\n{sub.text}")
    except Exception as e:  # noqa: BLE001
        return IngestResult(fmt, "FAILED", confidence="none", detail=f"archive error: {type(e).__name__}")
    return IngestResult(fmt, "EXTRACTED", "\n\n".join(texts), members, members=members,
                        detail=f"{len(members)} members")


def _ingest_docx(data: bytes) -> IngestResult:
    if not HAVE["docx"]:
        return IngestResult("docx", "BLOCKED", confidence="none", detail="[BLOCKED: install python-docx]")
    try:
        import docx
        d = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in d.paragraphs)
        return IngestResult("docx", "EXTRACTED", text, detail=f"{len(d.paragraphs)} paragraphs")
    except BaseException as e:  # noqa: BLE001
        return IngestResult("docx", "FAILED", confidence="none", detail=f"{type(e).__name__}")


def _ingest_xlsx(data: bytes) -> IngestResult:
    if not HAVE["xlsx"]:
        return IngestResult("xlsx", "BLOCKED", confidence="none", detail="[BLOCKED: install openpyxl]")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                rows.append([("" if c is None else str(c)) for c in row])
        return IngestResult("xlsx", "EXTRACTED", "\n".join(", ".join(r) for r in rows), rows,
                            detail=f"{len(rows)} rows")
    except BaseException as e:  # noqa: BLE001
        return IngestResult("xlsx", "FAILED", confidence="none", detail=f"{type(e).__name__}")


def _ingest_pdf(data: bytes) -> IngestResult:
    if not HAVE["pdf"]:
        return IngestResult("pdf", "BLOCKED", confidence="none",
                            detail="[BLOCKED: pypdf unavailable (broken cryptography backend in this env)]")
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
        if not text.strip():                         # no text layer ⇒ scanned ⇒ needs OCR (honest)
            return IngestResult("pdf", "EXTRACTED_LOWCONF", "", confidence="low",
                                detail="no text layer — scanned PDF needs OCR [BLOCKED: no tesseract]")
        return IngestResult("pdf", "EXTRACTED", text, detail=f"{len(reader.pages)} pages")
    except BaseException as e:  # noqa: BLE001
        return IngestResult("pdf", "FAILED", confidence="none", detail=f"{type(e).__name__}")


def _ingest_image(data: bytes) -> IngestResult:
    if not HAVE["ocr"]:
        return IngestResult("image", "BLOCKED", confidence="none",
                            detail="[BLOCKED: OCR engine (tesseract + pytesseract) not installed]")
    try:
        import pytesseract
        from PIL import Image
        text = pytesseract.image_to_string(Image.open(io.BytesIO(data)))
        conf = "high" if text.strip() else "low"
        return IngestResult("image", "EXTRACTED" if text.strip() else "EXTRACTED_LOWCONF", text,
                            confidence=conf, detail="OCR")
    except BaseException as e:  # noqa: BLE001
        return IngestResult("image", "FAILED", confidence="none", detail=f"{type(e).__name__}")


_ROUTES = {"json": _ingest_json, "ipynb": _ingest_ipynb, "csv": _ingest_csv,
           "docx": _ingest_docx, "xlsx": _ingest_xlsx, "pdf": _ingest_pdf, "image": _ingest_image}


def ingest(data: bytes, filename: str = "") -> IngestResult:
    """Detect the format and extract; route to S21 grounding via `.text`. Unknown/corrupt → honest FAILED."""
    fmt = detect_format(filename, data)
    try:
        if fmt in _ROUTES:
            return _ROUTES[fmt](data)
        if fmt in ("zip", "tar"):
            return _ingest_archive(data, fmt)
        if fmt in ("text", "code"):
            return _ingest_text(data, fmt)
        return IngestResult(fmt, "FAILED", confidence="none", detail="unsupported / undetected format")
    except Exception as e:  # noqa: BLE001
        return IngestResult(fmt, "FAILED", confidence="none", detail=f"extraction error: {type(e).__name__}")


def available_formats() -> Dict[str, bool]:
    """Honest report of what this environment can extract (optional libs/binaries detected at import)."""
    return {"json": True, "ipynb": True, "csv": True, "text": True, "code": True, "zip": True, "tar": True,
            "docx": HAVE["docx"], "xlsx": HAVE["xlsx"], "pdf": HAVE["pdf"], "image(OCR)": HAVE["ocr"]}
